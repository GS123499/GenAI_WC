import streamlit as st
import os
from openai import AzureOpenAI
import javaproperties
from docx import Document
from docx.shared import RGBColor
from datetime import datetime
from logger import log_message  
from docx.shared import RGBColor
from docx.shared import Pt, RGBColor
from docx.enum.style import WD_STYLE_TYPE
import tempfile
import base64
 
import re
import networkx as nx
from pyvis.network import Network
from collections import defaultdict
import glob
#print("Glob module loaded:", glob)
 
 
# Load Properties
def load_properties(path):
    with open(path, "r") as f:
        return javaproperties.load(f)
   
import os
import glob
 
def load_java_files(properties, use_config_files=True):
    files_to_process = []
 
    if use_config_files:
        for key, raw_path in properties.items():
            if key in ["subscription_key", "endpoint", "api_version", "deployment"]:
                continue
 
            # Convert to OS-compatible path (e.g., handle forward/back slashes)
            raw_path = os.path.normpath(raw_path)
 
            matched_files = []
 
            # If it's a .java file directly and it exists, take it
            if raw_path.endswith(".java") and os.path.isfile(raw_path):
                matched_files.append(raw_path)
 
            # If it's a directory, recursively fetch all .java files inside
            elif os.path.isdir(raw_path):
                matched_files.extend(glob.glob(os.path.join(raw_path, "**", "*.java"), recursive=True))
 
            # If it's a glob pattern (e.g., **/*.java), expand it
            elif "*" in raw_path or "?" in raw_path:
                matched_files.extend(glob.glob(raw_path, recursive=True))
 
            else:
                # Handle edge case: path doesn't exist or is unrecognized
                if os.path.exists(raw_path):
                    print(f"⚠️ Path exists but is not a file or recognized pattern: {raw_path}")
                else:
                    print(f"⚠️ Path does not exist: {raw_path}")
 
            if not matched_files:
                print(f"⚠️ No files matched for: {raw_path}")
 
            for file_path in matched_files:
                if file_path.endswith(".java") and os.path.isfile(file_path):
                    try:
                        with open(file_path, "r", encoding="utf-8", errors="replace") as f:
                            content = f.read()
                            files_to_process.append((os.path.basename(file_path), file_path, content))
                    except Exception as e:
                        print(f"❌ Error reading {file_path}: {e}")
 
    return files_to_process
 
 
properties = load_properties("resources/config.properties")
files = load_java_files(properties)
 
# Extract OpenAI Credentials
subscription_key = properties.get("subscription_key")
endpoint = properties.get("endpoint")
api_version = properties.get("api_version")
deployment_name = properties.get("deployment")
 
if not all([subscription_key, endpoint, api_version, deployment_name]):
    raise ValueError("Missing one or more OpenAI config values")
 
# Initialize Azure OpenAI Client
client = AzureOpenAI(
    api_key=subscription_key,
    azure_endpoint=endpoint,
    api_version=api_version,
)
 
def chunk_code(code, chunk_size=500):
    lines = code.splitlines()
    return ["\n".join(lines[i:i + chunk_size]) for i in range(0, len(lines), chunk_size)]
 
def get_openai_response(prompt):
    try:
        response = client.chat.completions.create(
            model=deployment_name,
            messages=[
                {"role": "system", "content": "You are a helpful assistant."},
                {"role": "user", "content": prompt}
            ],
            max_tokens=1000,
            temperature=0.3,
        )
        return response.choices[0].message.content.strip()
    except Exception as e:
        log_message(f"OpenAI error: {e}")
        return f"Error during explanation: {str(e)}"
 
def get_explanation(user_input, code_chunk, reference_content):
    prompt = f"""
I am a PTC windchill developer generate Technical design documentation including method signatures, purpose, and data flow. Use the provided reference to explain the code.
 
User Prompt:
{user_input}
 
Code:
{code_chunk}
 
Reference:
{reference_content if reference_content else 'No reference provided.'}
 
Provide clear explanation using ONLY the reference context.
"""
    return get_openai_response(prompt)
 
def extract_class_name(content):
    match = re.search(r'public\s+class\s+(\w+)', content)
    return match.group(1) if match else "UnknownClass"
 
def save_to_word(doc, file_name, explanation, error_flag):
    try:
        doc.add_heading(f"Class: {class_name}", level=1)
        if not error_flag:
            for section in explanation.strip().split("\n\n"):
                lines = section.strip().split("\n")
                if len(lines) > 1:
                    heading = lines[0].lstrip("#").strip()
                    body = "\n".join(lines[1:]).strip()
                    heading_para = doc.add_paragraph()
                    heading_run = heading_para.add_run(heading)
                    heading_run.font.size = Pt(11)
                    heading_run.bold = True
                    heading_run.font.color.rgb = RGBColor(0, 0, 0)
 
                    doc.add_paragraph(body)
                else:
                    doc.add_paragraph(section.strip())
        else:
            doc.add_paragraph(f"[{file_name}] Error occurred. Explanation could not be generated.")
 
    except Exception as e:
        print("Error while adding explanation to Word doc:", e)      
       
 
 
# --- KNOWLEDGE GRAPH FUNCTIONS ---
OOTB_METHODS = {
    "wt.part.WTPart": ["doOperation", "getName", "getNumber"],
    "wt.fc.Persistable": ["save", "delete"],
    "wt.util.WTException": ["toString"],
    "wt.method.RemoteAccess": ["remoteMethod"]
}
 
def extract_class_name_and_parent(content):
    match = re.search(r'public\s+class\s+(\w+)(?:\s+extends\s+([\w\.]+))?', content)
    return match.groups() if match else (None, None)
 
def extract_methods(content):
    return re.findall(r'public\s+(?:static\s+)?(?:[\w\<\>\[\]]+\s+)+(\w+)\s*\(', content)
 
def extract_method_calls(content):
    return re.findall(r'(\w+)\.(\w+)\s*\(', content)
 
def extract_imports(content):
    return re.findall(r'import\s+([\w\.]+);', content)
 
def shorten(name):
    return name if len(name) <= 20 else name[:17] + "..."
 
def build_graph_from_files(files):
    G = nx.DiGraph()
 
    class_methods = defaultdict(list)         # {ClassName: [method1, method2]}
    class_parents = {}                        # {ChildClass: ParentClass}
    class_imports = defaultdict(set)          # {ClassName: set(imports)}
    method_definitions = {}                   # {methodName: definingClass}
    method_calls = defaultdict(set)           # {methodName: set(callingClasses)}
    all_classes = set()                       # Track all class names
 
    OOTB_SIMPLE_MAP = {cls.split('.')[-1]: methods for cls, methods in OOTB_METHODS.items()}
 
    for fname, _, content in files:
        class_name, parent_class = extract_class_name_and_parent(content)
        methods = extract_methods(content)
        calls = extract_method_calls(content)
        imports = extract_imports(content)
 
        if class_name:
            all_classes.add(class_name)
            class_methods[class_name] = methods
            for m in methods:
                method_definitions[m] = class_name
            if parent_class:
                class_parents[class_name] = parent_class
            class_imports[class_name].update(imports)
 
        for _, method in calls:
            if class_name:
                method_calls[method].add(class_name)
 
    # --- Add class nodes ---
    for class_name in all_classes:
        G.add_node(class_name, label=shorten(class_name), color="mediumpurple", title=f"Class: {class_name}")
 
    # --- Add method nodes and DEFINES edges ---
    for class_name, methods in class_methods.items():
        for method in methods:
            G.add_node(method, label=shorten(method), color="orange", title=f"Method: {method}")
            G.add_edge(class_name, method, label="DEFINES", color="gray")
 
    # --- Add method call edges (CALLS) ---
    for method, callers in method_calls.items():
        if method in method_definitions:
            for caller in callers:
                G.add_edge(caller, method, label="CALLS", color="green")
 
    # --- Add EXTENDS edges ---
    for child, parent in class_parents.items():
        G.add_node(parent, label=shorten(parent), color="lightblue", title=f"Parent: {parent}")
        G.add_edge(child, parent, label="EXTENDS", color="blue")
 
    # --- Add IMPORTS edges ---
    for cls, imports in class_imports.items():
        for imp in imports:
            imported_class = imp.split('.')[-1]
            G.add_node(imported_class, label=shorten(imported_class), color="lightgray", title=f"Imported: {imp}")
            G.add_edge(cls, imported_class, label="IMPORTS", color="darkgray")
 
    # --- OOTB Overrides ---
    for class_name, parent in class_parents.items():
        methods = class_methods.get(class_name, [])
        ootb_methods = OOTB_SIMPLE_MAP.get(parent.split('.')[-1], [])
        for method in methods:
            if method in ootb_methods:
                ootb_class_node = parent
                ootb_method_node = f"{parent}.{method}"
 
                G.add_node(ootb_class_node, label=shorten(ootb_class_node), color="#ccf2ff", title="OOTB Class")
                G.add_node(ootb_method_node, label=method, color="#ff9999", title="OOTB Method")
                G.add_edge(ootb_class_node, ootb_method_node, label="DEFINES", color="#999")
                G.add_edge(method, ootb_method_node, label="OVERRIDES", color="red")
 
    return G
 
def render_graph_pyvis(G):
    net = Network(height="800px", width="100%", bgcolor="white", font_color="black", directed=True)
    for node, data in G.nodes(data=True):
        net.add_node(node, label=data.get("label", node), color=data.get("color", "lightgray"), title=data.get("title", node))
    for source, target, data in G.edges(data=True):
        net.add_edge(source, target, label=data.get("label", ""), arrows="to")
    net.force_atlas_2based(gravity=-30, spring_length=200)
    tmp_file = tempfile.NamedTemporaryFile(delete=False, suffix=".html")
    net.write_html(tmp_file.name)
    return tmp_file.name
# Streamlit UI
# Page config
st.set_page_config(page_title="Windchill Code Analyzer", layout="wide")
 
# Injected CSS
st.markdown("""
<style>
/* General compact layout */
.stApp {
    background-color: white;
    color: #222;
    font-size: 12px;
}
 
/* Input fields styling */
.stTextInput input, .stTextArea textarea {
    font-size: 12px !important;
    padding: 6px !important;
    height: 28px !important;
    border: none;
}
 
/* File uploader styles */
div[data-testid="stFileUploader"] {
    border: 1px solid #c084fc !important;
    border-radius: 6px !important;
    background-color: #f3e8ff !important;
    padding: 6px 8px !important;
    margin-bottom: 8px;
    max-width: 100%;
}
div[data-testid="stFileUploader"] section {
    background-color: #f3e8ff !important;
    border-radius: 4px !important;
    padding: 4px !important;
}
div[data-testid="stFileUploader"] label {
    font-size: 10px !important;
    margin-bottom: 2px;
    color: #3a0576;
}
div[data-testid="stFileUploader"] button {
    background-color: #6108c4 !important;
    color: white !important;
    font-size: 11px !important;
    font-weight: bold;
    border-radius: 5px !important;
    padding: 4px 10px !important;
    height: 28px !important;
    min-width: 100px !important;
    border: none !important;
}
div[data-testid="stFileUploader"] button:hover {
    background-color: #420585 !important;
    color: white !important;
    cursor: pointer;
}
 
/* Checkbox styling */
.css-1aumxhk {
    font-size: 12px !important;
}
 
/* Button styling */
.stButton button {
    font-size: 12px !important;
    padding: 4px 10px !important;
    height: 32px !important;
    border-radius: 6px;
    background-color: #6108c4 !important;
    font-weight: bold;
    color: white !important;
    border: none;
}
.stButton > button:hover {
    background-color: #420585 !important;
}
 
/* Title animation */
@keyframes fadeSlideIn {
    0% { opacity: 0; transform: translateY(-10px); }
    100% { opacity: 1; transform: translateY(0); }
}
 
/* Full page layout */
.block-container {
    padding-top: 3rem !important;
}
 
/* Header section */
.violet-header {
    background-color: #6A0DAD;
    padding: 20px;
    border-radius: 15px;
    margin-bottom: 20px;
    color: white;
    width: 100%;
    display: flex;
    align-items: center;
    justify-content: flex-start;
    gap: -40px; /* Increased space between logo and text */
}
.logo-container {
    flex-shrink: 0;
}
.header-logo {
    height: 45px;
    border-radius: 6px;
    margin-left: 20px; /* Push logo a bit to the right */
}
.header-text {
    text-align: center;
    flex-grow: 1;
}
.header-text h2 {
    margin: 0;
    font-size: 40px;
    font-weight: bold;
}
.header-text p {
    margin: -10px 0 0;
    font-size: 13px;
}
 
/* Query heading smaller */
.query-heading {
    font-size: 1.2rem;
    font-weight: bold;
    display: flex;
    align-items: center;
    gap: 0.5rem;
    margin-bottom: 0.5rem;
    margin-top: -1rem;
    color: #420585;
}
</style>
""", unsafe_allow_html=True)
 
# Logo image
with open("C:/Asset_GenAI/logoacc.png", "rb") as img_file:
    encoded = base64.b64encode(img_file.read()).decode()
 
# Header layout
st.markdown(f"""
<div class="violet-header">
    <img src="data:image/png;base64,{encoded}" class="header-logo" alt="Accenture Logo">
    <div class="header-text">
        <h2>Windchill Code Analyzer</h2>
        <p>Upload code and a reference document. Ask your query and get AI-based explanations using your context.</p>
    </div>
</div>
""", unsafe_allow_html=True)
 
# Query heading
st.markdown("""<div class="query-heading">💬 Enter Your Query</div>""", unsafe_allow_html=True)
user_prompt = st.text_input("Enter your query", placeholder="Type your question here...", label_visibility="collapsed")
 
# File uploaders side by side
col1, col2 = st.columns(2)
with col1:
    uploaded_file = st.file_uploader("📂 Upload Java File", type=["java"], accept_multiple_files=True)
    use_config_files = not uploaded_file
with col2:
    reference_file = st.file_uploader("📑 Upload Reference File", type=["pdf", "docx", "txt"])
 
# Centered checkbox and submit button
col1, col2, col3, col4, col5 = st.columns([1.7, 1.7, 2, 1.5, 1.5])
with col3:
    generate_graph = st.checkbox("📊 Generate Knowledge Graph")
   
btn_col1, btn_col2, btn_col3 = st.columns([0.7, 1, 2])
with btn_col3:
    submitted = st.button("🚀 Submit")
 
 
 
# Processing Logic
if submitted:
    if not user_prompt:
        st.warning("Please enter a query.")
        st.stop()
 
    # Initialize Word doc
    doc = Document()
    doc.add_heading("Code Explanation Report", 0)
    #doc.add_paragraph(f"Prompt: {user_prompt}")
    doc.add_paragraph(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
 
   
    # Read reference file
    reference_content = ""
    if reference_file:
        try:
            if reference_file.name.endswith(".docx"):
                docx_doc = Document(reference_file)
                reference_content = "\n".join([para.text for para in docx_doc.paragraphs])
            elif reference_file.name.endswith(".pdf"):
                import PyPDF2
                reader = PyPDF2.PdfReader(reference_file)
                reference_content = "\n".join([page.extract_text() for page in reader.pages if page.extract_text()])
            elif reference_file.name.endswith(".txt"):
                reference_content = reference_file.read().decode("utf-8")
        except Exception as e:
            st.error(f"Error reading reference file: {e}")
            st.stop()
 
    # --- Load Java files ---
    files_to_process = []
    if uploaded_file:
        for single_file in uploaded_file:
            code_content = single_file.read().decode("utf-8", errors="ignore")
            files_to_process.append((single_file.name, single_file.name, code_content))
    elif use_config_files:
        files_to_process = load_java_files(properties)
 
    if not files_to_process:
        st.warning("No Java files available for explanation. Upload a file or check config paths.")
        st.stop()
 
    # --- Process each Java file ---
    with st.spinner("⚙️ Processing files..."):
        for file_name, file_path, code in files_to_process:
        #st.write(f"🔄 Processing: `{file_name}`")
            chunks = chunk_code(code)
            combined_explanation = ""
            error_flag = False
 
        for i, chunk in enumerate(chunks, 1):
            explanation = get_explanation(user_prompt, chunk, reference_content)
            if "Error during explanation" in explanation:
                explanation = f"[Error in chunk {i}] {explanation}"
                error_flag = True
                log_message(explanation)
            combined_explanation += f"\n{explanation.strip()}\n"
        class_name = extract_class_name(code)
        save_to_word(doc, file_name, combined_explanation, error_flag)
       
    # --- Save Word document ---
    try:
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        final_output_path = f"Code_Explanation_Report_{timestamp}.docx"
        doc.save(final_output_path)
        st.success(f"✅ All explanations saved to: `{final_output_path}`")
    except PermissionError:
            st.error("❌ Please close the Word document before running again.")
           
        # --- Generate and Save Excel Summary ---
    try:
        import re
        import pandas as pd


        def extract_method_details(content):
            # Strict method declaration matcher
            method_pattern = re.compile(
                r'^\s*(public|private|protected)?\s*'
                r'(static\s+)?'
                r'([\w\<\>\[\]]+)\s+'
                r'(\w+)\s*'
                r'\(([^)]*)\)\s*'
                r'(\{)?',
                re.MULTILINE
            )

            param_pattern = re.compile(r'((?:final\s+)?[\w\<\>\[\]]+\s+\w+)')

            matches = method_pattern.findall(content)
            method_details = []

            for access, static, return_type, method_name, params, _ in matches:
                if params.strip():
                    param_list = [p.strip() for p in param_pattern.findall(params)]
                    formatted_params = ", ".join(param_list)
                else:
                    formatted_params = ""

                method_details.append({
                    "Access Modifier": access if access else "default",
                    "Static": "Yes" if static and static.strip() else "No",
                    "Return Type": return_type.strip(),
                    "Method Name": method_name.strip(),
                    "Parameters": formatted_params
                })

            return method_details

 
        summary_data = []
 
        for file_name, file_path, content in files_to_process:
            try:
                class_name, parent_class = extract_class_name_and_parent(content)
                method_info = extract_method_details(content)
 
                for method in method_info:
                    summary_data.append({
                        "Class Name": class_name or "Unknown",
                        "Parent Class": parent_class or "None",
                        "Access Modifier": method["Access Modifier"],
                        "Static": method["Static"],
                        "Return Type": method["Return Type"],
                        "Method Name": method["Method Name"],
                        "Parameters": method["Parameters"]
                    })
            except Exception as inner_err:
                st.warning(f"⚠️ Skipped file `{file_name}` due to parsing error: {inner_err}")
 
        if summary_data:
            df = pd.DataFrame(summary_data)
            excel_path = f"Code_Detailed_Summary_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx"
            df.to_excel(excel_path, index=False)
            st.success(f"📊 Detailed Excel summary saved to: `{excel_path}`")
        else:
            st.warning("⚠️ No valid method data found for Excel summary.")
 
    except Exception as e:
        st.error(f"❌ Failed to generate Excel summary: {e}")
       
 
    if generate_graph:
        st.subheader("📌 Dependency Graph")
        graph = build_graph_from_files(files_to_process)
        html_path = render_graph_pyvis(graph)
        with open(html_path, 'r', encoding='utf-8') as f:
            graph_html = f.read()
        st.components.v1.html(graph_html, height=850, scrolling=True)
 
